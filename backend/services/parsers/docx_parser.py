from pathlib import Path
from docx import Document
from backend.models.parsed_document import ParsedDocument


class DOCXParser:

    def parse(self, file_path: str) -> ParsedDocument:

        path = Path(file_path)

        document = Document(path)

        content = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

        return ParsedDocument(
            source=str(path),
            file_type="docx",
            content=content,
            metadata={
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "paragraphs": len(document.paragraphs)
            }
        )